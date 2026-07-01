#!/usr/bin/env python3
"""
efi_load_report.py - Analyse SDSEFI electrical load log files.

Reads the same two-column .dat files used by efi_plot.py and produces an HTML, PDF, or DOCX report.

Important implementation notes:
    Raw overview plots can be decimated for long logs, but the decimator must
    preserve the shape of narrow pulses. A simple t[::N], y[::N] decimator can
    skip the zero/near-zero samples between pulses and then draw a misleading
    line from one pulse directly to the next. This script uses an envelope
    decimator for injector and coil raw overview plots: each display bucket
    contributes its first point, minimum point, maximum point, and last point
    in time order. That keeps pulse returns-to-zero visible while still
    reducing plot size.

    Fuelpump runtime logs are often sparse/intermittent samples rather than a
    continuous waveform. Those traces are plotted as markers only so the report
    does not imply a straight-line current transition across long time gaps.

Expected files, when present:
  log_injector_0.dat ... log_injector_5.dat
  log_coil_left.dat, log_coil_right.dat
  log_fuelpump_left.dat, log_fuelpump_right.dat
  log_left_pumpstart_detail.dat, log_right_pumpstart_detail.dat

HTML dependencies:
  numpy, plotly

PDF dependencies:
  numpy, plotly, kaleido, reportlab

DOCX dependencies:
  numpy, plotly, kaleido, python-docx
"""

from __future__ import annotations

import argparse
import html
import io
import math
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional, Sequence, Tuple

import numpy as np
import plotly.graph_objects as go
from plotly.subplots import make_subplots


INJECTOR_FILES = [(f"Injector {i + 1}", f"log_injector_{i}.dat") for i in range(6)]
COIL_FILES = [("Left coilpack", "log_coil_left.dat"), ("Right coilpack", "log_coil_right.dat")]
FUELPUMP_FILES = [("Left fuelpump", "log_fuelpump_left.dat"), ("Right fuelpump", "log_fuelpump_right.dat")]
PUMPSTART_FILES = [("Left fuelpump start", "log_left_pumpstart_detail.dat"), ("Right fuelpump start", "log_right_pumpstart_detail.dat")]


@dataclass
class Channel:
    name: str
    filename: str
    kind: str
    t: np.ndarray
    y: np.ndarray


@dataclass
class PulseMetrics:
    channel: str
    count: int
    threshold: float
    starts: np.ndarray
    durations_ms: np.ndarray
    peaks: np.ndarray
    areas: np.ndarray
    slopes: np.ndarray

    def median(self, values: np.ndarray) -> float:
        return float(np.median(values)) if values.size else float("nan")

    def mean(self, values: np.ndarray) -> float:
        return float(np.mean(values)) if values.size else float("nan")

    def std(self, values: np.ndarray) -> float:
        return float(np.std(values)) if values.size else float("nan")




@dataclass
class PlotScope:
    offset: Optional[float] = None
    scope: Optional[float] = None

    @property
    def enabled(self) -> bool:
        return self.offset is not None or self.scope is not None

    def range_for(self, tmin: float, tmax: float) -> Optional[Tuple[float, float]]:
        if not self.enabled or not np.isfinite(tmin) or not np.isfinite(tmax):
            return None
        if self.offset is None:
            # If --scope is supplied without --offset, centre the window at
            # exactly halfway through the available plot data.
            start = (tmin + tmax) / 2.0
        else:
            start = self.offset
        if self.scope is None:
            end = tmax
        else:
            end = start + self.scope
        return (float(start), float(end))


def common_time_extent(channels: Sequence[Channel]) -> Optional[Tuple[float, float]]:
    mins = [float(np.nanmin(ch.t)) for ch in channels if ch.t.size]
    maxs = [float(np.nanmax(ch.t)) for ch in channels if ch.t.size]
    if not mins or not maxs:
        return None
    return min(mins), max(maxs)


def apply_x_scope(fig: go.Figure, scope: PlotScope, tmin: float, tmax: float) -> None:
    xr = scope.range_for(tmin, tmax)
    if xr is None:
        return
    fig.update_xaxes(range=[xr[0], xr[1]])

def load_two_column(path: Path) -> Optional[Tuple[np.ndarray, np.ndarray]]:
    """Load a two-column log file and return sorted, finite x/y arrays."""
    if not path.exists():
        return None
    try:
        arr = np.loadtxt(path, delimiter=",", ndmin=2)
    except ValueError:
        arr = np.loadtxt(path, ndmin=2)

    if arr.ndim != 2 or arr.shape[1] < 2 or arr.shape[0] == 0:
        return None

    arr = np.asarray(arr[:, :2], dtype=float)
    finite = np.isfinite(arr[:, 0]) & np.isfinite(arr[:, 1])
    arr = arr[finite]
    if arr.shape[0] == 0:
        return None

    # The logger normally writes monotonic time, but sorting makes the report
    # robust to concatenated or hand-edited files.
    order = np.argsort(arr[:, 0], kind="mergesort")
    arr = arr[order]
    return arr[:, 0], arr[:, 1]


def load_numeric_matrix(path: Path) -> Optional[np.ndarray]:
    """Load a whitespace- or comma-separated numeric matrix.

    Pump-start detail files produced by dl.c/gnuplot are not two-column logs.
    Column 1 is elapsed time in milliseconds and columns 2..N are separate
    pump-start captures.  Use genfromtxt rather than loadtxt so partially
    empty rows, if ever present, become NaN and individual records can still
    be used.
    """
    if not path.exists():
        return None

    for kwargs in ({"delimiter": None}, {"delimiter": ","}):
        try:
            arr = np.genfromtxt(
                path,
                ndmin=2,
                dtype=float,
                comments="#",
                invalid_raise=False,
                **kwargs,
            )
            if arr.size and arr.ndim == 2 and arr.shape[1] >= 2:
                return np.asarray(arr, dtype=float)
        except Exception:
            pass
    return None


def load_pumpstart_titles(data_dir: Path, data_filename: str) -> dict[int, str]:
    """Return optional gnuplot titles for pump-start data columns.

    dl.c may write lpump.gpl/rpump.gpl containing commands such as:
        'log_right_pumpstart_detail.dat' using 1:4 title "Pump start ..."
    These titles are useful labels, but the data file itself remains the
    source of truth.
    """
    import re

    titles: dict[int, str] = {}
    gpl_candidates = []
    if "right" in data_filename:
        gpl_candidates.append(data_dir / "rpump.gpl")
    if "left" in data_filename:
        gpl_candidates.append(data_dir / "lpump.gpl")
    gpl_candidates.extend(sorted(data_dir.glob("*pump.gpl")))

    pattern = re.compile(
        r"['\"](?P<file>[^'\"]+)['\"]\s+using\s+1:(?P<col>\d+)\s+title\s+['\"](?P<title>[^'\"]+)['\"]"
    )
    for gpl in gpl_candidates:
        if not gpl.exists():
            continue
        try:
            text = gpl.read_text(errors="replace")
        except OSError:
            continue
        for match in pattern.finditer(text):
            if Path(match.group("file")).name == data_filename:
                titles[int(match.group("col"))] = match.group("title")
    return titles


def load_pumpstart_channels(data_dir: Path) -> List[Channel]:
    """Load pump-start logs, one Channel per capture column.

    For pump-start files, column 1 is elapsed time in milliseconds and every
    later column is an independent pump-start record.  The normal two-column
    loader must not be used, otherwise only column 2 is ever plotted.
    """
    out: List[Channel] = []
    for base_name, fn in PUMPSTART_FILES:
        arr = load_numeric_matrix(data_dir / fn)
        if arr is None or arr.shape[1] < 2:
            continue

        t_all = arr[:, 0]
        titles = load_pumpstart_titles(data_dir, fn)
        record_count = arr.shape[1] - 1
        for col in range(1, arr.shape[1]):
            y_all = arr[:, col]
            finite = np.isfinite(t_all) & np.isfinite(y_all)
            if np.count_nonzero(finite) < 2:
                continue

            t = t_all[finite]
            y = y_all[finite]
            # Preserve file order; do not sort.  Pump-start time is already
            # elapsed time and each column is a complete record.
            if col + 1 in titles:
                title = titles[col + 1]
                # Keep names short enough for legends/tables while retaining
                # the useful timestamp if dl.c provided one.
                name = f"{base_name} {col}: {title}" if record_count > 1 else f"{base_name}: {title}"
            else:
                name = base_name if record_count == 1 else f"{base_name} {col}"

            out.append(Channel(name=name, filename=f"{fn}:col{col + 1}", kind="pumpstart", t=t, y=y))
    return out


def load_channels(data_dir: Path) -> Tuple[List[Channel], List[Channel], List[Channel], List[Channel]]:
    groups: List[List[Channel]] = [[], [], []]
    specs = [INJECTOR_FILES, COIL_FILES, FUELPUMP_FILES]
    kinds = ["injector", "coil", "fuelpump"]
    for group, spec, kind in zip(groups, specs, kinds):
        for name, fn in spec:
            loaded = load_two_column(data_dir / fn)
            if loaded is None:
                continue
            t, y = loaded
            group.append(Channel(name=name, filename=fn, kind=kind, t=t, y=y))

    pumpstarts = load_pumpstart_channels(data_dir)
    return groups[0], groups[1], groups[2], pumpstarts


def estimate_sample_period(t: np.ndarray) -> float:
    dt = np.diff(t)
    dt = dt[np.isfinite(dt) & (dt > 0)]
    if dt.size == 0:
        return 0.0
    return float(np.median(dt))


def decimate_envelope_xy(t: np.ndarray, y: np.ndarray, max_points: int) -> Tuple[np.ndarray, np.ndarray]:
    """Shape-preserving decimation for pulse trains.

    The old report script used simple stride decimation, equivalent to
    t[::step], y[::step]. That is unsafe for narrow pulse/current logs because
    the samples that prove a pulse has returned to zero can be skipped. The
    plotted line then appears to bridge from one pulse into the next.

    This function divides the trace into buckets and emits first, min, max and
    last points from each bucket, preserving their original time order. It also
    keeps both sides of unusually large time gaps, which makes sparse pulse logs
    display as honestly as possible.
    """
    n = len(t)
    if max_points <= 0 or n <= max_points or n < 3:
        return t, y

    # Four points per bucket are emitted, so target roughly max_points/4 buckets.
    bucket_size = max(2, int(math.ceil(4.0 * n / float(max_points))))
    keep: List[int] = []

    dt_nom = estimate_sample_period(t)
    gap_idx: set[int] = set()
    if dt_nom > 0:
        # Keep boundaries around acquisition/logging gaps. This is especially
        # useful for logs that only record windows around active pulses.
        gaps = np.where(np.diff(t) > max(5.0 * dt_nom, 0.001))[0]
        for i in gaps:
            gap_idx.add(int(i))
            gap_idx.add(int(i + 1))

    for start in range(0, n, bucket_size):
        end = min(start + bucket_size, n)
        if end <= start:
            continue
        seg = y[start:end]
        first = start
        last = end - 1
        imin = start + int(np.argmin(seg))
        imax = start + int(np.argmax(seg))
        # Preserve temporal order so Plotly draws the line correctly.
        keep.extend(sorted({first, imin, imax, last}))

    keep.extend(gap_idx)
    keep.extend([0, n - 1])
    idx = np.asarray(sorted(set(i for i in keep if 0 <= i < n)), dtype=int)
    return t[idx], y[idx]


def detect_pulses(ch: Channel, threshold: Optional[float] = None, min_width_s: float = 0.00015) -> PulseMetrics:
    t = ch.t
    y = ch.y
    empty = np.array([], dtype=float)
    if y.size == 0:
        return PulseMetrics(ch.name, 0, 0.0 if threshold is None else threshold, empty, empty, empty, empty, empty)

    ymax = float(np.max(y))
    if threshold is None:
        if ch.kind == "injector":
            threshold = max(0.05, 0.12 * ymax)
        elif ch.kind == "coil":
            threshold = max(0.10, 0.08 * ymax)
        else:
            threshold = max(0.05, 0.10 * ymax)

    mask = y > threshold
    changes = np.diff(mask.astype(np.int8))
    starts_idx = np.where(changes == 1)[0] + 1
    ends_idx = np.where(changes == -1)[0] + 1
    if mask[0]:
        starts_idx = np.r_[0, starts_idx]
    if mask[-1]:
        ends_idx = np.r_[ends_idx, len(mask)]

    starts: List[float] = []
    durations_ms: List[float] = []
    peaks: List[float] = []
    areas: List[float] = []
    slopes: List[float] = []

    for s, e in zip(starts_idx, ends_idx):
        if e <= s + 1:
            continue
        duration = float(t[e - 1] - t[s])
        if duration < min_width_s:
            continue
        seg_t = t[s:e]
        seg_y = y[s:e]
        peak = float(np.max(seg_y))
        area = float(np.trapezoid(seg_y, seg_t))
        pidx = int(np.argmax(seg_y))
        dt_peak = float(seg_t[pidx] - seg_t[0]) if pidx > 0 else float("nan")
        slope = (peak - float(seg_y[0])) / dt_peak if dt_peak and dt_peak > 0 else float("nan")
        starts.append(float(t[s]))
        durations_ms.append(duration * 1000.0)
        peaks.append(peak)
        areas.append(area)
        slopes.append(slope)

    return PulseMetrics(
        channel=ch.name,
        count=len(starts),
        threshold=float(threshold),
        starts=np.asarray(starts),
        durations_ms=np.asarray(durations_ms),
        peaks=np.asarray(peaks),
        areas=np.asarray(areas),
        slopes=np.asarray(slopes),
    )



def fig_overview(channels: Sequence[Channel], inj_scale: float, max_points: int, scope: PlotScope) -> go.Figure:
    rows = len(channels)
    fig = make_subplots(
        rows=rows,
        cols=1,
        shared_xaxes=True,
        vertical_spacing=0.012,
        subplot_titles=[c.name for c in channels],
    )
    for row, ch in enumerate(channels, start=1):
        tx, yy = decimate_envelope_xy(ch.t, ch.y, max_points)
        if ch.kind == "fuelpump":
            fig.add_trace(
                go.Scattergl(
                    x=tx,
                    y=yy,
                    mode="markers",
                    marker=dict(size=4),
                    name=ch.name,
                    showlegend=False,
                ),
                row=row,
                col=1,
            )
        else:
            fig.add_trace(go.Scattergl(x=tx, y=yy, mode="lines", name=ch.name, showlegend=False), row=row, col=1)
        if ch.kind == "injector":
            fig.update_yaxes(range=[0.0, inj_scale], row=row, col=1, title_text="A")
        elif ch.kind in ("coil", "fuelpump"):
            ymax = max(7.0, float(np.nanmax(ch.y)) * 1.1 if ch.y.size else 7.0)
            fig.update_yaxes(range=[0.0, ymax], row=row, col=1, title_text="A")
    fig.update_xaxes(title_text="Time (s)", row=rows, col=1)
    extent = common_time_extent(channels)
    if extent is not None:
        apply_x_scope(fig, scope, extent[0], extent[1])
    fig.update_layout(
        height=max(850, rows * 170),
        title="Raw electrical load overview",
        margin=dict(l=70, r=30, t=80, b=50),
        hovermode="x unified",
    )
    return fig


def fig_metric_bars(
    title: str,
    metrics: Sequence[PulseMetrics],
    field: str,
    ylabel: str,
    bar_width: Optional[float] = None,
) -> go.Figure:
    names = [m.channel for m in metrics]
    vals = []
    for m in metrics:
        arr = getattr(m, field)
        vals.append(float(np.median(arr)) if arr.size else 0.0)

    bar_kwargs = {"x": names, "y": vals}
    if bar_width is not None:
        bar_kwargs["width"] = bar_width

    fig = go.Figure(go.Bar(**bar_kwargs))
    fig.update_layout(
        title=title,
        yaxis_title=ylabel,
        xaxis_title="Channel",
        height=450,
        margin=dict(l=70, r=30, t=70, b=110),
    )
    return fig


def fig_injector_peak_width_bars(metrics: Sequence[PulseMetrics], bar_width: float = 0.16) -> go.Figure:
    """Combined injector peak-current and pulse-width comparison.

    The old report used two separate one-series bar plots.  With only six
    injector channels, Plotly's default categorical bar width makes each bar
    visually much wider than useful.  This figure combines the two related
    injector metrics and deliberately narrows the bars.
    """
    names = [m.channel for m in metrics]

    peak_vals: List[float] = []
    width_vals: List[float] = []
    for m in metrics:
        peak_vals.append(float(np.median(m.peaks)) if m.peaks.size else 0.0)
        width_vals.append(float(np.median(m.durations_ms)) if m.durations_ms.size else 0.0)

    fig = go.Figure()
    fig.add_trace(
        go.Bar(
            x=names,
            y=peak_vals,
            name="Median peak current",
            yaxis="y",
            width=bar_width,
            offsetgroup="peak",
        )
    )
    fig.add_trace(
        go.Bar(
            x=names,
            y=width_vals,
            name="Median pulse width",
            yaxis="y2",
            width=bar_width,
            offsetgroup="width",
        )
    )

    fig.update_layout(
        title="Injector median peak current and pulse width",
        xaxis_title="Channel",
        yaxis=dict(title="Median peak current (A)", side="left"),
        yaxis2=dict(title="Median pulse width (ms)", overlaying="y", side="right"),
        barmode="group",
        bargap=0.65,
        bargroupgap=0.25,
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1.0),
        height=450,
        margin=dict(l=75, r=85, t=90, b=110),
    )
    return fig


def fig_coilpack_peak_width_bars(metrics: Sequence[PulseMetrics], bar_width: float = 0.16) -> go.Figure:
    """Combined coilpack peak-current and dwell/current pulse-width comparison.

    This mirrors the injector combined plot: peak current uses the left Y axis,
    dwell/current pulse width uses the right Y axis, and the bars are kept
    deliberately narrow so two coilpack categories do not produce oversized
    blocks.
    """
    names = [m.channel for m in metrics]

    peak_vals: List[float] = []
    width_vals: List[float] = []
    for m in metrics:
        peak_vals.append(float(np.median(m.peaks)) if m.peaks.size else 0.0)
        width_vals.append(float(np.median(m.durations_ms)) if m.durations_ms.size else 0.0)

    fig = go.Figure()
    fig.add_trace(
        go.Bar(
            x=names,
            y=peak_vals,
            name="Median peak current",
            yaxis="y",
            width=bar_width,
            offsetgroup="peak",
        )
    )
    fig.add_trace(
        go.Bar(
            x=names,
            y=width_vals,
            name="Median dwell/current pulse width",
            yaxis="y2",
            width=bar_width,
            offsetgroup="width",
        )
    )

    fig.update_layout(
        title="Coilpack median peak current and dwell/current pulse width",
        xaxis_title="Channel",
        yaxis=dict(title="Median peak current (A)", side="left"),
        yaxis2=dict(title="Median dwell/current pulse width (ms)", overlaying="y", side="right"),
        barmode="group",
        bargap=0.65,
        bargroupgap=0.25,
        legend=dict(orientation="h", yanchor="bottom", y=1.02, xanchor="right", x=1.0),
        height=450,
        margin=dict(l=75, r=95, t=90, b=110),
    )
    return fig


def fig_event_raster(metrics: Sequence[PulseMetrics], scope: PlotScope) -> go.Figure:
    fig = go.Figure()
    for idx, m in enumerate(metrics):
        if m.starts.size == 0:
            continue
        fig.add_trace(
            go.Scattergl(
                x=m.starts,
                y=np.full_like(m.starts, idx, dtype=float),
                mode="markers",
                marker=dict(symbol="line-ns-open", size=10),
                name=m.channel,
            )
        )
    all_starts = np.concatenate([m.starts for m in metrics if m.starts.size]) if any(m.starts.size for m in metrics) else np.array([])
    if all_starts.size:
        apply_x_scope(fig, scope, float(np.nanmin(all_starts)), float(np.nanmax(all_starts)))
    fig.update_layout(
        title="Detected pulse timing",
        xaxis_title="Time (s)",
        yaxis=dict(tickmode="array", tickvals=list(range(len(metrics))), ticktext=[m.channel for m in metrics]),
        height=500,
        margin=dict(l=130, r=30, t=70, b=50),
    )
    return fig


def extract_pulse_windows(ch: Channel, m: PulseMetrics, pre_s: float, post_s: float, max_pulses: int) -> Tuple[np.ndarray, List[np.ndarray]]:
    if m.starts.size == 0:
        return np.array([]), []
    indices = np.linspace(0, len(m.starts) - 1, min(max_pulses, len(m.starts))).astype(int)
    rel = np.linspace(-pre_s, post_s, 220)
    windows: List[np.ndarray] = []
    for i in indices:
        centre = m.starts[i]
        x = centre + rel
        if x[0] < ch.t[0] or x[-1] > ch.t[-1]:
            continue
        windows.append(np.interp(x, ch.t, ch.y))
    return rel * 1000.0, windows


def fig_overlay(channels: Sequence[Channel], metrics: Sequence[PulseMetrics], title: str, pre_s: float, post_s: float, max_pulses: int) -> go.Figure:
    rows = len(channels)
    fig = make_subplots(rows=rows, cols=1, shared_xaxes=True, vertical_spacing=0.03, subplot_titles=[c.name for c in channels])
    for row, (ch, m) in enumerate(zip(channels, metrics), start=1):
        rel_ms, windows = extract_pulse_windows(ch, m, pre_s, post_s, max_pulses)
        for w in windows:
            fig.add_trace(go.Scattergl(x=rel_ms, y=w, mode="lines", opacity=0.25, showlegend=False), row=row, col=1)
        if windows:
            med = np.median(np.vstack(windows), axis=0)
            fig.add_trace(go.Scatter(x=rel_ms, y=med, mode="lines", name=f"{ch.name} median", line=dict(width=3), showlegend=False), row=row, col=1)
        fig.update_yaxes(title_text="A", row=row, col=1)
    fig.update_xaxes(title_text="Time relative to detected rising edge (ms)", row=rows, col=1)
    fig.update_layout(title=title, height=max(550, rows * 230), margin=dict(l=70, r=30, t=80, b=50))
    return fig


def fig_pulse_metric_trend(metrics: Sequence[PulseMetrics], title: str, field: str, ylabel: str, scope: PlotScope) -> go.Figure:
    fig = go.Figure()
    for m in metrics:
        arr = getattr(m, field)
        if m.starts.size and arr.size:
            fig.add_trace(go.Scattergl(x=m.starts, y=arr, mode="markers", name=m.channel))
    all_starts = np.concatenate([m.starts for m in metrics if m.starts.size]) if any(m.starts.size for m in metrics) else np.array([])
    if all_starts.size:
        apply_x_scope(fig, scope, float(np.nanmin(all_starts)), float(np.nanmax(all_starts)))
    fig.update_layout(title=title, xaxis_title="Time (s)", yaxis_title=ylabel, height=500, margin=dict(l=70, r=30, t=70, b=50))
    return fig


def moving_average(y: np.ndarray, requested_window: int) -> np.ndarray:
    if y.size < 3:
        return y
    win = max(3, int(requested_window))
    if win % 2 == 0:
        win += 1
    win = min(win, y.size if y.size % 2 == 1 else y.size - 1)
    if win < 3:
        return y
    kernel = np.ones(win) / win
    pad = win // 2
    padded = np.pad(y, (pad, pad), mode="edge")
    return np.convolve(padded, kernel, mode="valid")


def fig_pumpstart(channels: Sequence[Channel], max_points: int, title: str) -> go.Figure:
    """Create one pump-start overlay plot for one pump side.

    Each Channel in *channels* represents one data column from a pump-start
    detail file: column 1 is elapsed time, and columns 2..N are separate
    captured starts.  Keep every capture as its own line, but do not add peak
    markers/text because they become unreadable with multiple overlaid starts.
    """
    fig = go.Figure()
    xlabel = "Time"
    for ch in channels:
        x = ch.t
        xlabel = "Time (ms)" if x.size and np.nanmax(x) > 20.0 else "Time (s)"
        tx, yy = decimate_envelope_xy(x, ch.y, max_points)
        fig.add_trace(go.Scattergl(x=tx, y=yy, mode="lines", name=ch.name))
    fig.update_layout(
        title=title,
        xaxis_title=xlabel,
        yaxis_title="Current (A)",
        height=500,
        margin=dict(l=70, r=30, t=70, b=50),
    )
    return fig


def split_pumpstart_sides(pumpstarts: Sequence[Channel]) -> Tuple[List[Channel], List[Channel]]:
    """Split pump-start capture channels into left and right pump groups."""
    left: List[Channel] = []
    right: List[Channel] = []
    for ch in pumpstarts:
        key = (ch.filename + " " + ch.name).lower()
        if "right" in key or "rpump" in key:
            right.append(ch)
        elif "left" in key or "lpump" in key:
            left.append(ch)
    return left, right


def summary_rows(metrics: Sequence[PulseMetrics]) -> List[List[str]]:
    rows = []
    for m in metrics:
        rows.append([
            m.channel,
            str(m.count),
            f"{m.threshold:.3f}",
            f"{m.median(m.peaks):.3f}",
            f"{m.std(m.peaks):.3f}",
            f"{m.median(m.durations_ms):.3f}",
            f"{m.std(m.durations_ms):.3f}",
            f"{m.median(m.areas) * 1000.0:.3f}",
        ])
    return rows


def html_table(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> str:
    out = ["<table>", "<thead><tr>" + "".join(f"<th>{html.escape(h)}</th>" for h in headers) + "</tr></thead>", "<tbody>"]
    for row in rows:
        out.append("<tr>" + "".join(f"<td>{html.escape(str(c))}</td>" for c in row) + "</tr>")
    out.extend(["</tbody>", "</table>"])
    return "\n".join(out)


def channel_stats_rows(channels: Sequence[Channel]) -> List[List[str]]:
    rows = []
    for ch in channels:
        if ch.kind in ("fuelpump", "pumpstart"):
            duration_text = ""
        else:
            duration = float(ch.t[-1] - ch.t[0]) if ch.t.size > 1 else 0.0
            duration_text = f"{duration:.3f}"

        if ch.kind == "pumpstart":
            sample_rate_text = ""
        else:
            dt = np.diff(ch.t)
            dtpos = dt[np.isfinite(dt) & (dt > 0)]
            fs = 1.0 / float(np.median(dtpos)) if dtpos.size and np.median(dtpos) > 0 else float("nan")
            sample_rate_text = f"{fs:.1f}" if np.isfinite(fs) else ""

        rows.append([
            ch.name,
            ch.filename,
            ch.kind,
            str(len(ch.t)),
            duration_text,
            sample_rate_text,
            f"{np.nanmax(ch.y):.3f}",
            f"{np.nanmean(ch.y):.3f}",
        ])
    return rows


def fig_to_html(fig: go.Figure) -> str:
    return fig.to_html(full_html=False, include_plotlyjs=False)


def fig_to_png_bytes(fig: go.Figure, width: int = 1200, height: Optional[int] = None) -> bytes:
    if height is None:
        height = int(fig.layout.height) if fig.layout.height else 650

    # Static export via Kaleido is much more reliable with ordinary scatter
    # traces than with WebGL scattergl traces, especially for DOCX/PDF output.
    # Keep Scattergl for interactive HTML, but convert a copy before rasterising.
    spec = fig.to_plotly_json()
    for trace in spec.get("data", []):
        if trace.get("type") == "scattergl":
            trace["type"] = "scatter"
    fig_static = go.Figure(spec)
    return fig_static.to_image(format="png", width=width, height=height, scale=1.3)


def write_html(path: Path, title: str, data_dir: Path, channels: Sequence[Channel], metrics: Sequence[PulseMetrics], figs: Sequence[go.Figure]) -> None:
    headers = ["Channel", "Pulses", "Threshold A", "Median peak A", "Peak SD", "Median width ms", "Width SD", "Median area mA.s"]
    stats_headers = ["Channel", "File", "Kind", "Samples", "Duration s", "Sample rate Hz", "Max A", "Mean A"]
    parts = [
        "<!doctype html><html><head><meta charset='utf-8'>",
        f"<title>{html.escape(title)}</title>",
        "<script src='https://cdn.plot.ly/plotly-2.35.2.min.js'></script>",
        "<style>body{font-family:Arial,sans-serif;margin:28px;line-height:1.35} table{border-collapse:collapse;margin:1em 0;width:100%;font-size:13px} th,td{border:1px solid #ccc;padding:5px 7px;text-align:right} th:first-child,th:nth-child(2),td:first-child,td:nth-child(2),td:nth-child(3){text-align:left} th:nth-child(3),td:nth-child(3){text-align:right} h1,h2{margin-top:1.4em}.note{background:#f5f5f5;padding:10px;border-left:4px solid #999}</style>",
        "</head><body>",
        f"<h1>{html.escape(title)}</h1>",
        f"<p><b>Input directory:</b> {html.escape(str(data_dir))}</p>",
        "<h2>Input data summary</h2>",
        html_table(stats_headers, channel_stats_rows(channels)),
        "<h2>Detected pulse metrics</h2>",
        html_table(headers, summary_rows(metrics)),
    ]
    for fig in figs:
        parts.append(fig_to_html(fig))
    parts.append("</body></html>")
    path.write_text("\n".join(parts), encoding="utf-8")


def write_pdf(path: Path, title: str, data_dir: Path, channels: Sequence[Channel], metrics: Sequence[PulseMetrics], figs: Sequence[go.Figure], landscape_pages: bool = False) -> None:
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib.units import cm
        from reportlab.platypus import Image, PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
    except ImportError as exc:
        raise SystemExit("PDF output requires reportlab. Install it with: pip install reportlab") from exc

    try:
        _ = fig_to_png_bytes(figs[0], width=1000, height=600) if figs else b""
    except Exception as exc:
        raise SystemExit("PDF output requires kaleido for Plotly image export. Install it with: pip install kaleido") from exc

    page_size = landscape(A4) if landscape_pages else A4
    doc = SimpleDocTemplate(str(path), pagesize=page_size, rightMargin=1.0 * cm, leftMargin=1.0 * cm, topMargin=1.0 * cm, bottomMargin=1.0 * cm)
    styles = getSampleStyleSheet()
    story = [
        Paragraph(title, styles["Title"]),
        Paragraph(f"Input directory: {html.escape(str(data_dir))}", styles["Normal"]),
        Spacer(1, 0.3 * cm),
    ]
    story.append(Paragraph("Input data summary", styles["Heading2"]))
    stat_data = [["Channel", "File", "Kind", "Samples", "Duration s", "Sample rate Hz", "Max A", "Mean A"]] + channel_stats_rows(channels)
    tbl = Table(stat_data, repeatRows=1)
    tbl.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 7)]))
    story.extend([tbl, Spacer(1, 0.3 * cm), Paragraph("Detected pulse metrics", styles["Heading2"])])
    metric_data = [["Channel", "Pulses", "Threshold A", "Med peak A", "Peak SD", "Med width ms", "Width SD", "Med area mA.s"]] + summary_rows(metrics)
    mtbl = Table(metric_data, repeatRows=1)
    mtbl.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), colors.lightgrey), ("GRID", (0, 0), (-1, -1), 0.25, colors.grey), ("FONTSIZE", (0, 0), (-1, -1), 7)]))
    story.extend([mtbl, PageBreak()])

    for fig in figs:
        png = fig_to_png_bytes(fig, width=1200)
        bio = io.BytesIO(png)
        img = Image(bio)
        max_w = 27.0 * cm
        max_h = 18.0 * cm
        scale = min(max_w / img.imageWidth, max_h / img.imageHeight)
        img.drawWidth = img.imageWidth * scale
        img.drawHeight = img.imageHeight * scale
        story.append(img)
        story.append(PageBreak())
    doc.build(story)


def write_docx(path: Path, title: str, data_dir: Path, channels: Sequence[Channel], metrics: Sequence[PulseMetrics], figs: Sequence[go.Figure], landscape_pages: bool = False) -> None:
    """Write a DOCX version of the same report content used for HTML/PDF.

    This intentionally uses python-docx rather than ODT/odfpy because LibreOffice
    Writer imports DOCX tables and embedded PNG figures reliably. Plotly figures
    are rendered to PNG using Kaleido, the same mechanism used by the PDF output.
    """
    try:
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
    except ImportError as exc:
        raise SystemExit("DOCX output requires python-docx. Install it with: pip install python-docx") from exc

    # Ensure Kaleido/static image export is available before creating the DOCX.
    try:
        _ = fig_to_png_bytes(figs[0], width=1000, height=600) if figs else b""
    except Exception as exc:
        raise SystemExit("DOCX output requires kaleido for Plotly image export. Install it with: pip install kaleido") from exc

    doc = Document()

    # Default to portrait A4-style pages. Use --landscape for wide PDF/DOCX output.
    section = doc.sections[0]
    if landscape_pages:
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.69)
        section.page_height = Inches(8.27)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Inches(8.27)
        section.page_height = Inches(11.69)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(9)

    def add_note(text: str) -> None:
        p = doc.add_paragraph(text)
        p.style = styles["Normal"]

    def add_table(headers: Sequence[str], rows: Sequence[Sequence[str]]) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(headers):
            hdr[i].text = str(h)
            for para in hdr[i].paragraphs:
                for run in para.runs:
                    run.bold = True
        for row in rows:
            cells = table.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
        doc.add_paragraph()

    def figure_title(fig: go.Figure, fallback: str) -> str:
        try:
            text = fig.layout.title.text
            return str(text) if text else fallback
        except Exception:
            return fallback

    def add_figure(fig: go.Figure, idx: int) -> None:
        title_text = figure_title(fig, f"Figure {idx}")
        doc.add_page_break()
        doc.add_heading(title_text, level=2)

        # Match the PDF path: render a static PNG and fit it to the landscape page.
        pixel_w = 1600
        pixel_h = int(fig.layout.height) if fig.layout.height else 700
        png = fig_to_png_bytes(fig, width=pixel_w, height=pixel_h)
        image_stream = io.BytesIO(png)

        max_w_in = 10.6 if landscape_pages else 7.25
        max_h_in = 6.55 if landscape_pages else 9.1
        scale = min(max_w_in / float(pixel_w), max_h_in / float(pixel_h))
        width_in = pixel_w * scale
        height_in = pixel_h * scale

        pic = doc.add_picture(image_stream)
        pic.width = Inches(width_in)
        pic.height = Inches(height_in)
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading(title, level=1)
    add_note(f"Input directory: {data_dir}")

    doc.add_heading("Input data summary", level=2)
    add_table(
        ["Channel", "File", "Kind", "Samples", "Duration s", "Sample rate Hz", "Max A", "Mean A"],
        channel_stats_rows(channels),
    )

    doc.add_heading("Detected pulse metrics", level=2)
    add_table(
        ["Channel", "Pulses", "Threshold A", "Med peak A", "Peak SD", "Med width ms", "Width SD", "Med area mA.s"],
        summary_rows(metrics),
    )

    for idx, fig in enumerate(figs, start=1):
        add_figure(fig, idx)

    doc.save(path)


def build_report(data_dir: Path, output: Path, fmt: str, inj_scale: float, max_plot_points: int, max_overlay_pulses: int, scope: PlotScope, landscape_pages: bool = False) -> None:
    injectors, coils, fuelpumps, pumpstarts = load_channels(data_dir)
    all_channels = injectors + coils + fuelpumps + pumpstarts
    if not all_channels:
        raise SystemExit(f"No recognised .dat files found in {data_dir}")
    if not pumpstarts:
        print("warning: no pump start detail file found; startup detail plot will be omitted")

    injector_metrics = [detect_pulses(ch, min_width_s=0.00025) for ch in injectors]
    coil_metrics = [detect_pulses(ch, min_width_s=0.00025) for ch in coils]
    all_metrics = injector_metrics + coil_metrics

    figs: List[go.Figure] = []
    figs.append(fig_overview(injectors + coils + fuelpumps, inj_scale=inj_scale, max_points=max_plot_points, scope=scope))
    if injector_metrics:
        figs.append(fig_event_raster(injector_metrics + coil_metrics, scope=scope))
        figs.append(fig_injector_peak_width_bars(injector_metrics, bar_width=0.16))
        figs.append(fig_metric_bars("Injector median current integral", injector_metrics, "areas", "Area (A.s)", bar_width=0.16))
        figs.append(fig_overlay(injectors, injector_metrics, "Injector aligned pulse overlays", pre_s=0.001, post_s=0.006, max_pulses=max_overlay_pulses))
        figs.append(fig_pulse_metric_trend(injector_metrics, "Injector peak current trend", "peaks", "Peak current (A)", scope=scope))
    if coil_metrics:
        figs.append(fig_coilpack_peak_width_bars(coil_metrics, bar_width=0.16))
        figs.append(fig_overlay(coils, coil_metrics, "Coilpack aligned current overlays", pre_s=0.001, post_s=0.012, max_pulses=max_overlay_pulses))
        figs.append(fig_pulse_metric_trend(coil_metrics, "Coilpack peak current trend", "peaks", "Peak current (A)", scope=scope))
    if pumpstarts:
        left_pumpstarts, right_pumpstarts = split_pumpstart_sides(pumpstarts)
        if left_pumpstarts:
            figs.append(fig_pumpstart(left_pumpstarts, max_points=max_plot_points, title="Left fuel pump startup detail"))
        if right_pumpstarts:
            figs.append(fig_pumpstart(right_pumpstarts, max_points=max_plot_points, title="Right fuel pump startup detail"))

    title = "SDSEFI electrical load analysis report"
    if fmt == "html":
        write_html(output, title, data_dir, all_channels, all_metrics, figs)
    elif fmt == "pdf":
        write_pdf(output, title, data_dir, all_channels, all_metrics, figs, landscape_pages=landscape_pages)
    elif fmt == "docx":
        write_docx(output, title, data_dir, all_channels, all_metrics, figs, landscape_pages=landscape_pages)
    else:
        raise SystemExit(f"Unsupported format: {fmt}")


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Analyse EFI electrical load .dat files and produce an HTML, PDF, or DOCX report.")
    p.add_argument("data_dir", nargs="?", default=".", help="Directory containing log_*.dat files")
    p.add_argument("-o", "--output", default=None, help="Output report filename. Default: efi_load_report.<format>")
    p.add_argument("--format", choices=["html", "pdf", "docx"], default="html", help="Report format to write (default: html)")
    p.add_argument("-i", "--inj-scale", type=float, default=1.0, metavar="MAX", help="Maximum Y-axis value for injector raw overview plots, in amps (default: 1.0)")
    p.add_argument("--max-plot-points", type=int, default=20000, help="Approximate maximum points per trace in raw trend plots (default: 20000). Envelope decimation may emit slightly more points to preserve pulse shape.")
    p.add_argument("--max-overlay-pulses", type=int, default=80, help="Maximum detected pulses to overlay per channel (default: 80)")
    p.add_argument("--offset", type=float, default=None, metavar="SECONDS", help="Start time, in seconds, for zoomed non-pumpstart plots. If omitted while --scope is supplied, the window starts halfway through the plot data.")
    p.add_argument("--scope", type=float, default=None, metavar="SECONDS", help="Duration, in seconds, of the zoomed non-pumpstart plot window. If omitted while --offset is supplied, plots run from --offset to the end of the data.")
    p.add_argument("--landscape", action="store_true", help="Use landscape page orientation for PDF and DOCX output. Has no effect for HTML.")
    return p.parse_args()


def main() -> None:
    args = parse_args()
    data_dir = Path(args.data_dir).expanduser().resolve()
    if not data_dir.is_dir():
        raise SystemExit(f"Input directory does not exist: {data_dir}")
    if args.inj_scale <= 0:
        raise SystemExit("--inj-scale must be positive")
    if args.max_plot_points < 1000:
        raise SystemExit("--max-plot-points must be at least 1000")
    if args.scope is not None and args.scope <= 0:
        raise SystemExit("--scope must be positive")
    output = Path(args.output) if args.output else Path(f"efi_load_report.{args.format}")
    build_report(
        data_dir=data_dir,
        output=output,
        fmt=args.format,
        inj_scale=args.inj_scale,
        max_plot_points=args.max_plot_points,
        max_overlay_pulses=args.max_overlay_pulses,
        scope=PlotScope(offset=args.offset, scope=args.scope),
        landscape_pages=args.landscape,
    )
    print(f"  Wrote {output}")


if __name__ == "__main__":
    main()
